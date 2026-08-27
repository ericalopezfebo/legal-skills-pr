#!/usr/bin/env python3
"""DOCX renderer for pr-motion-drafting.

Formatting helper only. It does not research law or decide legal sufficiency.
Requires python-docx. YAML input additionally requires PyYAML.
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    import yaml  # type: ignore
except Exception:
    yaml = None

FORUMS = {
    "casp": {
        "header": ["GOBIERNO DE PUERTO RICO", "COMISIÓN APELATIVA DEL SERVICIO PÚBLICO"],
        "salutation": "A LA HONORABLE COMISIÓN:", "case_label": "Caso Núm.", "matter_label": "SOBRE:",
    },
    "aep-ja": {
        "header": ["AUTORIDAD DE EDIFICIOS PÚBLICOS", "JUNTA DE APELACIONES", "SAN JUAN, PUERTO RICO"],
        "salutation": "A LA HONORABLE JUNTA:", "case_label": "Caso Núm.", "matter_label": "SOBRE:",
    },
    "tpi": {
        "header": ["ESTADO LIBRE ASOCIADO DE PUERTO RICO", "TRIBUNAL DE PRIMERA INSTANCIA"],
        "salutation": "AL HONORABLE TRIBUNAL:", "case_label": "Civil Núm.", "matter_label": "SOBRE:",
    },
    "admin-generic": {
        "header": ["GOBIERNO DE PUERTO RICO"],
        "salutation": "AL HONORABLE FORO:", "case_label": "Caso Núm.", "matter_label": "SOBRE:",
    },
}


def set_cell_borders(cell, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attrs in edges.items():
        tag = "w:" + edge_name
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn("w:" + key), str(value))


def set_page_borders(section, left=None, right=None) -> None:
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgBorders"))
    if pg is None:
        pg = OxmlElement("w:pgBorders")
        sect_pr.append(pg)
    for name, attrs in (("left", left), ("right", right)):
        old = pg.find(qn("w:" + name))
        if old is not None:
            pg.remove(old)
        if attrs:
            node = OxmlElement("w:" + name)
            for key, value in attrs.items():
                node.set(qn("w:" + key), str(value))
            pg.append(node)


def set_default_font(doc, font_name: str, size_pt: float) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    for key in ("ascii", "hAnsi", "cs"):
        style._element.rPr.rFonts.set(qn("w:" + key), font_name)


def add_run(p, text: str, bold=False) -> None:
    r = p.add_run(text)
    r.bold = bold


def add_header(doc, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_run(p, line, True)


def add_caption(doc, data: dict[str, Any], profile: dict[str, Any]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    left.width, right.width = Inches(3.6), Inches(2.9)
    left.vertical_alignment = right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for cell in (left, right):
        set_cell_borders(cell, top={"val":"nil"}, bottom={"val":"nil"}, left={"val":"nil"}, right={"val":"nil"})
    set_cell_borders(left,
        bottom={"val":"double","sz":"6","space":"0","color":"auto"},
        right={"val":"double","sz":"6","space":"0","color":"auto"})
    set_cell_borders(right, left={"val":"double","sz":"6","space":"0","color":"auto"})

    parties = data.get("parties", {}).get("left", [])
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, party in enumerate(parties):
        if i:
            p = left.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(party.get("name", "[PARTE]")).upper(), True)
        role = party.get("role")
        if role:
            q = left.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_run(q, str(role))
        if i == 0 and len(parties) > 1:
            q = left.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_run(q, "v.", True)

    p = right.paragraphs[0]
    add_run(p, f"{profile['case_label']} {data.get('case_number','[POR COMPLETAR]')}", True)
    if data.get("room"):
        q = right.add_paragraph(); add_run(q, f"Sala: {data['room']}")
    q = right.add_paragraph(); add_run(q, profile["matter_label"], True)
    q = right.add_paragraph(); add_run(q, str(data.get("matter", "[POR COMPLETAR]")))


def add_centered_bold(doc, text: str) -> None:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    add_run(p, text, True)


def add_body(doc, text: str) -> None:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15; p.paragraph_format.space_after = Pt(6)
    add_run(p, text)


def add_page_number(section) -> None:
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("- ")
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld); p.add_run(" -")


def build_motion(data: dict[str, Any]) -> Document:
    profile = FORUMS.get(data.get("forum_profile", "admin-generic"), FORUMS["admin-generic"])
    visual = data.get("visual_profile", "official-neutral")
    doc = Document(); sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)

    if visual == "pr-litigation-redline":
        sec.left_margin, sec.right_margin = Inches(1.5), Inches(0.5)
        sec.top_margin, sec.bottom_margin = Inches(0.625), Inches(0.3125)
        set_page_borders(sec,
            left={"val":"double","sz":"12","space":"4","color":"FF0000"},
            right={"val":"single","sz":"12","space":"4","color":"FF0000"})
    else:
        sec.left_margin = Inches(float(data.get("left_margin", 1.0)))
        sec.right_margin = Inches(float(data.get("right_margin", 1.0)))
        sec.top_margin = Inches(float(data.get("top_margin", 1.0)))
        sec.bottom_margin = Inches(float(data.get("bottom_margin", 1.0)))
    set_default_font(doc, data.get("font", "Times New Roman"), float(data.get("font_size", 12)))

    add_header(doc, profile["header"] + list(data.get("extra_header", [])))
    add_caption(doc, data, profile)
    add_centered_bold(doc, str(data.get("title", "MOCIÓN")).upper())

    p = doc.add_paragraph(); add_run(p, str(data.get("salutation", profile["salutation"])), True)
    if data.get("appearance"):
        add_body(doc, str(data["appearance"]))

    for block in data.get("sections", []):
        if block.get("heading"):
            add_centered_bold(doc, str(block["heading"]))
        for paragraph in block.get("paragraphs", []):
            add_body(doc, str(paragraph))
    for paragraph in data.get("body", []):
        add_body(doc, str(paragraph))

    if data.get("prayer"):
        text = str(data["prayer"])
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        leads = ("POR TODO LO CUAL,", "POR LO EXPUESTO,", "EN MÉRITO DE LO ANTERIOR,")
        matched = next((x for x in leads if text.upper().startswith(x)), None)
        if matched:
            add_run(p, text[:len(matched)], True); add_run(p, text[len(matched):])
        else:
            add_run(p, text)

    if data.get("respectfully_submitted", True):
        p = doc.add_paragraph(); add_run(p, str(data.get("respectfully_text", "RESPETUOSAMENTE SOMETIDO.")), True)

    if data.get("certification"):
        text = str(data["certification"]); p = doc.add_paragraph()
        if text.upper().startswith("CERTIFICO"):
            lead = text.split(":", 1)[0] + (":" if ":" in text else "")
            add_run(p, lead, True); add_run(p, text[len(lead):])
        else:
            add_run(p, text)

    if data.get("place_date"):
        add_body(doc, str(data["place_date"]))

    sig = data.get("signature") or {}
    for key in ("firm", "name", "rua", "address", "phone", "email"):
        if sig.get(key):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
            add_run(p, str(sig[key]), key in {"firm", "name"})

    if data.get("page_numbers", True):
        add_page_number(sec)
    return doc


def load_data(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8")
    if path.suffix.lower() == ".json":
        return json.loads(text)
    if path.suffix.lower() in {".yaml", ".yml"}:
        if yaml is None:
            raise SystemExit("PyYAML required for YAML input; use JSON or install pyyaml")
        return yaml.safe_load(text)
    try:
        return json.loads(text)
    except Exception:
        if yaml is None:
            raise SystemExit("Input is not JSON and PyYAML is unavailable")
        return yaml.safe_load(text)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("input", type=Path)
    ap.add_argument("output", type=Path)
    args = ap.parse_args()
    doc = build_motion(load_data(args.input))
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)


if __name__ == "__main__":
    main()
